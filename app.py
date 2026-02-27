import streamlit as st
import requests
import hashlib
import time
import os
import tempfile
from pathlib import Path
from urllib.parse import urlencode, urlparse, parse_qs

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Course Assistant",
    page_icon="📚",
    layout="centered",
    initial_sidebar_state="collapsed",
)

# ── Styling ────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Serif+Display:ital@0;1&family=DM+Sans:wght@300;400;500;600&display=swap');

html, body, [class*="css"] { font-family: 'DM Sans', sans-serif; }
.stApp { background: #0f1117; }
#MainMenu, footer, header { visibility: hidden; }
.block-container { padding-top: 1.5rem; padding-bottom: 1rem; max-width: 780px; }

.hero {
    text-align: center;
    padding: 1.5rem 1rem 1rem;
    border-bottom: 1px solid #1e2330;
    margin-bottom: 1rem;
}
.hero h1 {
    font-family: 'DM Serif Display', serif;
    font-size: 1.9rem;
    color: #e8ecf4;
    margin: 0 0 0.3rem;
    letter-spacing: -0.02em;
}
.hero .subtitle { font-size: 0.82rem; color: #6b7694; font-weight: 300; letter-spacing: 0.04em; text-transform: uppercase; }
.hero .badge {
    display: inline-block;
    background: #1a2540;
    border: 1px solid #2a3a60;
    color: #7a9fd4;
    font-size: 0.72rem;
    padding: 0.2rem 0.7rem;
    border-radius: 20px;
    margin-top: 0.6rem;
    font-weight: 500;
}

[data-testid="stChatMessageContent"] {
    background: #161b2e !important;
    border: 1px solid #1e2744 !important;
    border-radius: 12px !important;
    color: #cdd5e8 !important;
    font-size: 0.9rem !important;
    line-height: 1.65 !important;
    padding: 0.85rem 1rem !important;
}
[data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatarUser"]) [data-testid="stChatMessageContent"] {
    background: #1a2d4a !important;
    border-color: #2a4470 !important;
    color: #dde4f5 !important;
}

.source-box {
    background: #0d1420;
    border: 1px solid #1a2744;
    border-left: 3px solid #3a6bc4;
    border-radius: 8px;
    padding: 0.5rem 0.85rem;
    margin-top: 0.6rem;
    font-size: 0.75rem;
    color: #5a7ab4;
    font-style: italic;
}

.status-bar {
    text-align: center;
    font-size: 0.72rem;
    color: #3a4a6a;
    margin-bottom: 0.75rem;
    letter-spacing: 0.03em;
}
.status-dot { display: inline-block; width: 6px; height: 6px; background: #2d8a4e; border-radius: 50%; margin-right: 5px; vertical-align: middle; }
.status-warn { display: inline-block; width: 6px; height: 6px; background: #c8922a; border-radius: 50%; margin-right: 5px; vertical-align: middle; }

.stChatInputContainer { background: #161b2e !important; border: 1px solid #1e2744 !important; border-radius: 12px !important; }
.stChatInputContainer textarea { color: #cdd5e8 !important; font-family: 'DM Sans', sans-serif !important; font-size: 0.88rem !important; }
.stSpinner > div { border-top-color: #3a6bc4 !important; }
</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# BRIGHTSPACE OAUTH + API HELPERS
# ══════════════════════════════════════════════════════════════════════════════

INSTANCE_URL  = "https://nbcctest.brightspace.com"
CLIENT_ID     = "2b9cbd14-1e83-4c45-beee-ac2d7f71ef84"
TOKEN_URL     = f"{INSTANCE_URL}/d2l/auth/api/token"
API_BASE      = f"{INSTANCE_URL}/d2l/api/le/1.67"


def get_access_token() -> str:
    """
    Exchange client credentials for a Bearer token using the OAuth 2.0
    client_credentials grant. Token is cached in session state and
    refreshed automatically when it expires.
    """
    now = time.time()
    if (
        "oauth_token" in st.session_state
        and st.session_state.get("oauth_expires_at", 0) > now + 60
    ):
        return st.session_state["oauth_token"]

    resp = requests.post(TOKEN_URL, data={
        "grant_type":    "client_credentials",
        "client_id":     CLIENT_ID,
        "client_secret": st.secrets["BS_CLIENT_SECRET"],
        "scope":         "content:file:read content:modules:read content:topics:read",
    }, timeout=15)

    if resp.status_code != 200:
        raise RuntimeError(
            f"Token request failed ({resp.status_code}): {resp.text[:300]}"
        )

    data = resp.json()
    st.session_state["oauth_token"]      = data["access_token"]
    st.session_state["oauth_expires_at"] = now + data.get("expires_in", 3600)
    return data["access_token"]


def api_get(path: str, token: str, stream: bool = False):
    """GET from the Brightspace LE API with Bearer auth."""
    resp = requests.get(
        f"{API_BASE}{path}",
        headers={"Authorization": f"Bearer {token}"},
        timeout=30,
        stream=stream,
    )
    if resp.status_code == 401:
        # Force token refresh on next call
        st.session_state.pop("oauth_token", None)
        raise RuntimeError("Brightspace token expired mid-session — please reload.")
    resp.raise_for_status()
    return resp


def fetch_course_name(org_unit_id: str, token: str) -> str:
    """Get the course name from the org units API."""
    try:
        resp = requests.get(
            f"{INSTANCE_URL}/d2l/api/lp/1.9/orgstructure/{org_unit_id}",
            headers={"Authorization": f"Bearer {token}"},
            timeout=15,
        )
        if resp.ok:
            return resp.json().get("Name", f"Course {org_unit_id}")
    except Exception:
        pass
    return f"Course {org_unit_id}"


def fetch_content_topics(org_unit_id: str, token: str) -> list[dict]:
    """
    Walk the course table of contents and return all file-type topics,
    excluding assessments (quizzes, dropbox, checklist, survey).
    """
    EXCLUDED_TYPES = {
        "quiz", "dropbox", "checklist", "survey",
        "selfassess", "scorm", "lti",
    }

    resp = api_get(f"/{org_unit_id}/content/toc", token)
    toc  = resp.json()

    topics = []

    def walk(modules):
        for mod in modules:
            for topic in mod.get("Topics", []):
                t_type = (topic.get("TypeIdentifier") or "").lower()
                if t_type in EXCLUDED_TYPES:
                    continue
                # Only include file/document-type topics
                if topic.get("TopicType") in (1,):   # 1 = File
                    topics.append({
                        "id":    topic["Id"],
                        "title": topic.get("Title", f"Topic {topic['Id']}"),
                    })
            walk(mod.get("Modules", []))

    walk(toc.get("Modules", []))
    return topics


def download_topic_text(org_unit_id: str, topic_id: int, token: str) -> str | None:
    """
    Download a topic file and extract its text.
    Supports: .txt, .html/.htm, .pdf, .docx
    Returns None if the file type is unsupported or download fails.
    """
    try:
        resp = api_get(f"/{org_unit_id}/content/topics/{topic_id}/file", token, stream=True)
        content_type = resp.headers.get("Content-Type", "")
        content_disp = resp.headers.get("Content-Disposition", "")

        # Determine extension
        ext = ""
        if "filename=" in content_disp:
            fname = content_disp.split("filename=")[-1].strip().strip('"')
            ext   = Path(fname).suffix.lower()
        elif "html" in content_type:
            ext = ".html"
        elif "pdf" in content_type:
            ext = ".pdf"
        elif "plain" in content_type:
            ext = ".txt"

        raw = resp.content

        if ext in (".txt",):
            return raw.decode("utf-8", errors="ignore")

        elif ext in (".html", ".htm"):
            from bs4 import BeautifulSoup
            return BeautifulSoup(raw, "html.parser").get_text(separator="\n")

        elif ext == ".pdf":
            import io
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(raw))
            return "\n".join(p.extract_text() or "" for p in reader.pages)

        elif ext == ".docx":
            import io
            import docx
            doc = docx.Document(io.BytesIO(raw))
            return "\n".join(p.text for p in doc.paragraphs)

        else:
            # Fallback: try to decode as text
            try:
                return raw.decode("utf-8", errors="ignore")
            except Exception:
                return None

    except Exception:
        return None


# ══════════════════════════════════════════════════════════════════════════════
# RAG PIPELINE
# ══════════════════════════════════════════════════════════════════════════════

@st.cache_resource(show_spinner=False)
def build_pipeline(org_unit_id: str, content_hash: str):
    """
    Build the RAG pipeline for a given course.
    Keyed by org_unit_id + content_hash so it rebuilds only when content changes.
    """
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_huggingface import HuggingFaceEmbeddings
    from langchain_chroma import Chroma
    from langchain_google_genai import ChatGoogleGenerativeAI
    from langchain_core.prompts import ChatPromptTemplate
    from langchain_core.output_parsers import StrOutputParser
    from langchain_core.runnables import RunnablePassthrough
    from langchain.schema import Document

    # The raw texts are passed in via session state (fetched before this call)
    raw_docs = st.session_state.get("raw_docs", [])
    if not raw_docs:
        raise ValueError("No course content could be retrieved from Brightspace.")

    # Build LangChain Documents
    documents = [
        Document(page_content=d["text"], metadata={"title": d["title"]})
        for d in raw_docs
        if d.get("text")
    ]

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=400, chunk_overlap=60,
        separators=["\n\n", "\n", ". ", " "],
    )
    chunks = splitter.split_documents(documents)

    embeddings = HuggingFaceEmbeddings(
        model_name="all-MiniLM-L6-v2",
        model_kwargs={"device": "cpu"},
    )
    vectorstore = Chroma.from_documents(chunks, embeddings)
    retriever   = vectorstore.as_retriever(search_kwargs={"k": 4})

    llm = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",
        google_api_key=st.secrets["GOOGLE_API_KEY"],
        max_output_tokens=1024,
    )

    course_name = st.session_state.get("course_name", "this course")
    prompt = ChatPromptTemplate.from_template(
        f"""You are a helpful course assistant for {course_name}.
Answer the student's question using ONLY the course content provided below.
Be concise, friendly, and precise.
If the answer is not in the provided content, say clearly:
"I can't find that in the course materials — please check with your instructor."
Do NOT make up information.

Course content:
{{context}}

Student question: {{question}}

Answer:"""
    )

    def format_docs(docs):
        return "\n\n".join(
            f"[{d.metadata.get('title','Untitled')}]\n{d.page_content}" for d in docs
        )

    chain = (
        {"context": retriever | format_docs, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )

    return chain, retriever


def fetch_and_cache_content(org_unit_id: str) -> str:
    """
    Fetch all course content via API, cache in session state.
    Returns a hash string used to key the pipeline cache.
    """
    token  = get_access_token()
    topics = fetch_content_topics(org_unit_id, token)

    raw_docs = []
    for topic in topics:
        text = download_topic_text(org_unit_id, topic["id"], token)
        if text and text.strip():
            raw_docs.append({"title": topic["title"], "text": text.strip()})

    st.session_state["raw_docs"]    = raw_docs
    st.session_state["course_name"] = fetch_course_name(org_unit_id, token)
    st.session_state["topic_count"] = len(raw_docs)

    # Hash the content so pipeline rebuilds only when files change
    combined = "".join(d["text"] for d in raw_docs)
    return hashlib.md5(combined.encode()).hexdigest()


def ask(chain_and_retriever, question: str) -> tuple[str, list]:
    chain, retriever = chain_and_retriever
    answer  = chain.invoke(question).strip()
    sources = retriever.invoke(question)
    return answer, sources


# ══════════════════════════════════════════════════════════════════════════════
# COURSE ID FROM URL PARAMS
# ══════════════════════════════════════════════════════════════════════════════

params      = st.query_params
org_unit_id = params.get("course_id", "297671")   # default to demo course


# ══════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════════════════════════════════════

if "messages" not in st.session_state:
    st.session_state.messages = []
if "pipeline_ready" not in st.session_state:
    st.session_state.pipeline_ready = False
if "last_org_unit" not in st.session_state:
    st.session_state.last_org_unit = None


# ══════════════════════════════════════════════════════════════════════════════
# HERO
# ══════════════════════════════════════════════════════════════════════════════

course_display = st.session_state.get("course_name", f"Course {org_unit_id}")
st.markdown(f"""
<div class="hero">
    <h1>📚 Course Assistant</h1>
    <div class="subtitle">{course_display}</div>
    <div class="badge">Powered by Gemini · Brightspace RAG</div>
</div>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# BUILD / REFRESH PIPELINE
# ══════════════════════════════════════════════════════════════════════════════

# Rebuild if course changed or not yet built
needs_build = (
    not st.session_state.pipeline_ready
    or st.session_state.last_org_unit != org_unit_id
)

if needs_build:
    with st.spinner("Connecting to Brightspace and indexing course content…"):
        try:
            content_hash = fetch_and_cache_content(org_unit_id)
            pipeline     = build_pipeline(org_unit_id, content_hash)
            st.session_state["pipeline"]      = pipeline
            st.session_state["pipeline_ready"] = True
            st.session_state["last_org_unit"]  = org_unit_id
            st.session_state.messages          = []   # clear chat on course switch
        except Exception as e:
            st.error(f"⚠️ Could not build pipeline: {e}")
            st.stop()

topic_count = st.session_state.get("topic_count", 0)
if topic_count > 0:
    st.markdown(f"""
    <div class="status-bar">
        <span class="status-dot"></span>
        {topic_count} content file{"s" if topic_count != 1 else ""} indexed from Brightspace
    </div>
    """, unsafe_allow_html=True)
else:
    st.markdown("""
    <div class="status-bar">
        <span class="status-warn"></span>
        No content files found in this course — add files in Brightspace first
    </div>
    """, unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# CHAT HISTORY
# ══════════════════════════════════════════════════════════════════════════════

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])
        if msg.get("source"):
            st.markdown(
                f'<div class="source-box">📄 {msg["source"]}</div>',
                unsafe_allow_html=True,
            )


# ══════════════════════════════════════════════════════════════════════════════
# CHAT INPUT
# ══════════════════════════════════════════════════════════════════════════════

if question := st.chat_input("Ask a question about this course…"):
    st.session_state.messages.append({"role": "user", "content": question})
    with st.chat_message("user"):
        st.markdown(question)

    with st.chat_message("assistant"):
        with st.spinner("Searching course content…"):
            try:
                answer, sources = ask(st.session_state["pipeline"], question)
                source_label = None
                if sources:
                    title = sources[0].metadata.get("title", "Course material")
                    snippet = sources[0].page_content[:80].replace("\n", " ").strip()
                    source_label = f"{title} — {snippet}…"
            except Exception as e:
                answer       = f"Sorry, something went wrong: {e}"
                source_label = None

        st.markdown(answer)
        if source_label:
            st.markdown(
                f'<div class="source-box">📄 {source_label}</div>',
                unsafe_allow_html=True,
            )

    st.session_state.messages.append({
        "role": "assistant", "content": answer, "source": source_label,
    })


# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════

if st.session_state.messages:
    st.markdown("""
    <div style="text-align:center;margin-top:1.5rem;font-size:0.7rem;color:#2a3a5a;">
        Answers are grounded in course content only · Always verify with your instructor
    </div>
    """, unsafe_allow_html=True)
